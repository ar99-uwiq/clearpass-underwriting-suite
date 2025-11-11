from docx import Document

def memo_to_docx(memo_text: str, out_path: str):
    doc = Document()
    doc.add_heading("Underwriting Memo", level=1)
    for para in memo_text.split("\n\n"):
        doc.add_paragraph(para)
    doc.save(out_path)
